"""
Resume parsing service.

Extracts raw text from uploaded resume files (PDF/DOCX), cleans it, and
performs a lightweight, rule-based detection of common resume sections
(Education, Experience, Skills, Projects, Certifications).

Design notes for future AI resume screening:
- This module raises its OWN exception hierarchy (`ResumeParsingError` and
  subclasses) rather than `fastapi.HTTPException`. That keeps it usable
  from contexts that aren't an HTTP request — background workers, batch
  re-parsing jobs, CLI tools, or a future AI scoring pipeline. The API
  layer is responsible for translating these into HTTP responses/logs.
- `parse_resume()` returns a single `ParsedResume` value object carrying
  everything downstream consumers (including a future NLP/AI stage)
  would want: cleaned text, word/char/page counts, a scanned-document
  flag, and non-fatal warnings — not just the two fields the current API
  happens to expose.
- This module intentionally does NOT do any AI-based extraction, skill
  parsing, or job-matching — that's a separate concern for later.
"""

from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from zipfile import BadZipFile

import fitz  # PyMuPDF
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.models.resume import ResumeFileType

# --------------------------------------------------------------------------
# Exceptions
# --------------------------------------------------------------------------
# Framework-agnostic on purpose — see module docstring.


class ResumeParsingError(Exception):
    """Base class for all resume parsing failures."""


class ResumeFileNotFoundError(ResumeParsingError):
    """The file referenced by the resume record doesn't exist on disk."""


class UnsupportedFileTypeError(ResumeParsingError):
    """The resume's file_type isn't one this parser knows how to handle."""


class CorruptedFileError(ResumeParsingError):
    """The file exists but couldn't be opened/read as a valid PDF/DOCX."""


class EncryptedFileError(ResumeParsingError):
    """The PDF is password-protected and couldn't be opened."""


class EmptyDocumentError(ResumeParsingError):
    """The file is zero bytes."""


class FileTooLargeError(ResumeParsingError):
    """The file exceeds the size limit this parser is willing to process."""


# --------------------------------------------------------------------------
# Config
# --------------------------------------------------------------------------

PREVIEW_LENGTH = 1000
SECTION_CONTENT_PREVIEW_LENGTH = 500

# Defense-in-depth limits. The upload endpoint already enforces a max
# upload size, but this parser may eventually be called from other
# entrypoints (re-parse job, AI pipeline), so it guards itself too.
MAX_PARSE_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_PDF_PAGES = 25
MAX_TEXT_LENGTH = 500_000

# Below this many characters of cleaned text, we treat the document as
# effectively textless (e.g. a scanned/image-based PDF with no OCR layer).
MIN_TEXT_LENGTH_THRESHOLD = 30

# Canonical section name -> regex patterns that identify its header line.
# Patterns are matched against a *normalized* line (see
# `_normalize_header_candidate`), so things like "== SKILLS ==", "1. Education:"
# and "**Projects**" all reduce to something patterns can match cleanly.
SECTION_PATTERNS: dict[str, list[str]] = {
    "Education": [
        r"^education(al)?\s*(background)?$",
        r"^academic\s+(background|qualifications?)$",
        r"^educational\s+qualifications?$",
        r"^qualifications?$",
    ],
    "Experience": [
        r"^(work|professional|relevant|career)?\s*experience$",
        r"^work\s+history$",
        r"^employment\s+history$",
        r"^career\s+history$",
        r"^internships?(\s+experience)?$",
    ],
    "Skills": [
        r"^(technical\s+|key\s+|core\s+)?skills?$",
        r"^skill\s*set$",
        r"^core\s+competenc(y|ies)$",
        r"^technical\s+proficienc(y|ies)$",
        r"^areas?\s+of\s+expertise$",
    ],
    "Projects": [
        r"^(key\s+|notable\s+|personal\s+|academic\s+|course\s+)?projects?$",
    ],
    "Certifications": [
        r"^certifications?(\s*(&|and)\s*licenses?)?$",
        r"^licenses?\s*(&|and)?\s*certifications?$",
        r"^(professional\s+)?training(s)?\s*(&|and)?\s*certifications?$",
        r"^certifications?\s*(&|and)?\s*training(s)?$",
    ],
}

_COMPILED_PATTERNS: dict[str, list[re.Pattern]] = {
    name: [re.compile(p, re.IGNORECASE) for p in patterns]
    for name, patterns in SECTION_PATTERNS.items()
}

# A line is only treated as a possible header if it's short — real section
# headers are single words/phrases, not full sentences.
_MAX_HEADER_LINE_LENGTH = 40


@dataclass
class ParsedResume:
    full_text: str
    preview: str
    detected_sections: dict[str, str | None]
    word_count: int
    char_count: int
    page_count: int | None  # None for DOCX, where "pages" isn't a stable concept
    is_scanned: bool
    warnings: list[str] = field(default_factory=list)


# --------------------------------------------------------------------------
# Text cleaning
# --------------------------------------------------------------------------

_CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
# Re-joins words split by a hyphen at a PDF line wrap, e.g. "develop-\nment".
_HYPHEN_LINEBREAK_RE = re.compile(r"(\w)-\n(\w)")
_MULTI_BLANK_LINES_RE = re.compile(r"\n{3,}")
_MULTI_SPACE_RE = re.compile(r"[ \t]{2,}")
_BULLET_CHARS = "•◦▪‣∙●○■□➤➢★✓✔"
_BULLET_LINE_RE = re.compile(rf"^[{re.escape(_BULLET_CHARS)}]\s*", re.MULTILINE)


def clean_text(raw_text: str) -> str:
    """
    Normalizes extracted resume text so section detection — and any future
    NLP/AI processing — works on consistent input instead of raw PDF/DOCX
    extraction artifacts.

    - Unicode-normalizes (NFKC) so visually-identical characters compare equal.
    - Strips non-printable control characters.
    - Re-joins words hyphenated across a PDF line break.
    - Normalizes assorted bullet glyphs to a plain "- " prefix.
    - Collapses repeated spaces/tabs and excessive blank lines.
    """
    if not raw_text:
        return ""

    text = unicodedata.normalize("NFKC", raw_text)
    text = _CONTROL_CHARS_RE.sub("", text)
    text = _HYPHEN_LINEBREAK_RE.sub(r"\1\2", text)
    text = _BULLET_LINE_RE.sub("- ", text)
    text = _MULTI_SPACE_RE.sub(" ", text)
    text = "\n".join(line.strip() for line in text.splitlines())
    text = _MULTI_BLANK_LINES_RE.sub("\n\n", text)
    return text.strip()


# --------------------------------------------------------------------------
# Text extraction
# --------------------------------------------------------------------------

def _extract_text_pdf(file_path: Path, warnings: list[str]) -> tuple[str, int]:
    """
    Extracts text from a PDF using PyMuPDF.

    Handles encrypted PDFs (raises EncryptedFileError unless an empty
    password opens them) and caps the number of pages read at
    MAX_PDF_PAGES, recording a warning rather than failing outright.

    Returns (text, total_page_count).
    """
    try:
        doc = fitz.open(str(file_path))
    except Exception as exc:  # noqa: BLE001
        raise CorruptedFileError(f"Failed to open PDF: {exc}") from exc

    try:
        if doc.is_encrypted and not doc.authenticate(""):
            raise EncryptedFileError(
                "PDF is password-protected and cannot be parsed."
            )

        total_pages = doc.page_count
        if total_pages == 0:
            raise CorruptedFileError("PDF contains no pages.")

        pages_to_read = min(total_pages, MAX_PDF_PAGES)
        if total_pages > MAX_PDF_PAGES:
            warnings.append(
                f"PDF has {total_pages} pages; only the first {MAX_PDF_PAGES} "
                f"were parsed."
            )

        text_parts: list[str] = []
        for page_index in range(pages_to_read):
            try:
                text_parts.append(doc[page_index].get_text())
            except Exception as exc:  # noqa: BLE001
                warnings.append(f"Failed to extract page {page_index + 1}: {exc}")

        return "\n".join(text_parts), total_pages
    except (EncryptedFileError, CorruptedFileError):
        raise
    except Exception as exc:  # noqa: BLE001
        raise CorruptedFileError(f"Failed to extract text from PDF: {exc}") from exc
    finally:
        doc.close()


def _extract_text_docx(file_path: Path) -> str:
    """Extracts all text (paragraphs + table cells) from a DOCX file."""
    try:
        document = Document(str(file_path))
    except PackageNotFoundError as exc:
        raise CorruptedFileError(
            f"File is not a valid Word document: {exc}"
        ) from exc
    except BadZipFile as exc:
        raise CorruptedFileError(
            f"DOCX file is corrupted (invalid zip container): {exc}"
        ) from exc
    except Exception as exc:  # noqa: BLE001
        raise CorruptedFileError(f"Failed to open DOCX: {exc}") from exc

    try:
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        raise CorruptedFileError(f"Failed to extract text from DOCX: {exc}") from exc


def extract_text(file_path: Path, file_type: ResumeFileType) -> str:
    """
    Convenience helper returning raw extracted text only, with no
    cleaning or metadata. Kept public for callers (tests, scripts, a
    future AI pipeline) that just want the text.
    """
    if file_type == ResumeFileType.pdf:
        text, _ = _extract_text_pdf(file_path, warnings=[])
        return text
    if file_type == ResumeFileType.docx:
        return _extract_text_docx(file_path)

    raise UnsupportedFileTypeError(f"Unsupported file type for parsing: {file_type}")


# --------------------------------------------------------------------------
# Section detection
# --------------------------------------------------------------------------

def _normalize_header_candidate(line: str) -> str:
    """
    Strips decorative wrapping so headers like "== SKILLS ==", "**Projects**",
    "1. Education:" and "### Experience" all reduce to their bare text
    before pattern matching.
    """
    stripped = line.strip()
    stripped = re.sub(r"^[^A-Za-z]+", "", stripped)
    stripped = re.sub(r"[^A-Za-z]+$", "", stripped)
    stripped = re.sub(r"\s+", " ", stripped)
    return stripped.strip()


def _match_section_header(line: str) -> str | None:
    """Returns the canonical section name if `line` looks like a header."""
    candidate = _normalize_header_candidate(line)
    if not candidate or len(candidate) > _MAX_HEADER_LINE_LENGTH:
        return None

    for canonical_name, patterns in _COMPILED_PATTERNS.items():
        for pattern in patterns:
            if pattern.match(candidate):
                return canonical_name
    return None


def detect_sections(text: str) -> dict[str, str | None]:
    """
    Scans cleaned text line-by-line for known section headers and captures
    the content that follows each one, up until the next recognized header
    (or end of document). If a section header repeats (e.g. two
    "Experience" blocks), their content is concatenated in order.

    Returns a dict with all five target section names as keys. The value
    is the captured content (truncated) if the section was found, or None
    if it wasn't detected.
    """
    detected: dict[str, str | None] = {name: None for name in SECTION_PATTERNS}

    current_section: str | None = None
    buffers: dict[str, list[str]] = {name: [] for name in SECTION_PATTERNS}

    for line in text.splitlines():
        header_match = _match_section_header(line)
        if header_match:
            current_section = header_match
            continue

        if current_section:
            stripped = line.strip()
            if stripped:
                buffers[current_section].append(stripped)

    for name, content_lines in buffers.items():
        if content_lines:
            content = " ".join(content_lines)
            detected[name] = content[:SECTION_CONTENT_PREVIEW_LENGTH]

    return detected


# --------------------------------------------------------------------------
# Public entrypoint
# --------------------------------------------------------------------------

def parse_resume(file_path: Path, file_type: ResumeFileType) -> ParsedResume:
    """
    Extracts, cleans, and analyzes a stored resume file.

    Args:
        file_path: absolute path to the resume file on disk.
        file_type: ResumeFileType.pdf or ResumeFileType.docx.

    Returns:
        ParsedResume with cleaned text, a 1000-char preview, detected
        sections, word/char/page counts, a scanned-document flag, and
        any non-fatal warnings encountered.

    Raises:
        ResumeFileNotFoundError: the file doesn't exist on disk.
        EmptyDocumentError: the file is 0 bytes.
        FileTooLargeError: the file exceeds MAX_PARSE_FILE_SIZE_BYTES.
        UnsupportedFileTypeError: file_type isn't pdf/docx.
        CorruptedFileError: the file couldn't be opened/read.
        EncryptedFileError: a PDF is password-protected.
    """
    if not file_path.exists():
        raise ResumeFileNotFoundError(f"Resume file not found on disk: {file_path}")

    file_size = file_path.stat().st_size
    if file_size == 0:
        raise EmptyDocumentError("Resume file is empty (0 bytes).")
    if file_size > MAX_PARSE_FILE_SIZE_BYTES:
        raise FileTooLargeError(
            f"Resume file is {file_size} bytes, which exceeds the "
            f"{MAX_PARSE_FILE_SIZE_BYTES}-byte limit for parsing."
        )

    warnings: list[str] = []

    if file_type == ResumeFileType.pdf:
        raw_text, page_count = _extract_text_pdf(file_path, warnings)
    elif file_type == ResumeFileType.docx:
        raw_text, page_count = _extract_text_docx(file_path), None
    else:
        raise UnsupportedFileTypeError(f"Unsupported file type for parsing: {file_type}")

    cleaned_text = clean_text(raw_text)

    if len(cleaned_text) > MAX_TEXT_LENGTH:
        cleaned_text = cleaned_text[:MAX_TEXT_LENGTH]
        warnings.append(
            f"Extracted text truncated to {MAX_TEXT_LENGTH} characters."
        )

    is_scanned = len(cleaned_text.strip()) < MIN_TEXT_LENGTH_THRESHOLD
    if is_scanned:
        warnings.append(
            "Little to no extractable text found. This document may be "
            "scanned/image-based and require OCR."
        )

    sections = detect_sections(cleaned_text)
    if not is_scanned and not any(sections.values()):
        warnings.append(
            "No recognizable section headers were found. The resume may "
            "use non-standard formatting."
        )

    return ParsedResume(
        full_text=cleaned_text,
        preview=cleaned_text[:PREVIEW_LENGTH],
        detected_sections=sections,
        word_count=len(cleaned_text.split()),
        char_count=len(cleaned_text),
        page_count=page_count,
        is_scanned=is_scanned,
        warnings=warnings,
    )
